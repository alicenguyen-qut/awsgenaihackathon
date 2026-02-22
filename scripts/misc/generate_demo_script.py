from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('🍳 MealBuddy — 3-Minute Demo Script', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('AWS GenAI Hackathon  |  Team 54')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True

doc.add_paragraph()

sections = [
    ("HOOK (0:00 – 0:15)", [
        ("screen", "[Black screen. One stat fades in.]"),
        ("voice", '"11 million people die every year from diet-related disease."'),
        ("voice", '"Not because they don\'t care. Because eating well was never designed to be easy."'),
    ]),
    ("THE PROBLEM (0:15 – 0:35)", [
        ("screen", "[Split screen: 3 panels fade in one by one]"),
        ("voice", '"Eating healthy today means juggling three separate problems."'),
        ("voice", '"What do I eat? — 20 minutes Googling recipes that don\'t match your diet."'),
        ("screen", "[Panel 1: Google search results, recipe blogs]"),
        ("voice", '"Did I buy everything? — You\'re at the supermarket. The list is at home."'),
        ("screen", "[Panel 2: empty fridge, forgotten grocery list]"),
        ("voice", '"Am I on track? — You have no idea if today\'s meals fit your goals."'),
        ("screen", "[Panel 3: calorie counting spreadsheet]"),
        ("voice", '"Three apps. Three tabs. Zero coordination. A dietitian costs $200 a session."'),
        ("screen", "[All 3 panels collapse into a question mark]"),
        ("voice", '"There had to be a better way."'),
    ]),
    ("THE SOLUTION (0:35 – 0:50)", [
        ("screen", "[MealBuddy logo appears]"),
        ("voice", '"Meet MealBuddy — one AI conversation that replaces all three."'),
        ("voice", '"It plans your meals, builds your shopping list, tracks your nutrition, and learns your preferences — powered by Amazon Bedrock."'),
        ("screen", "[Persona card: Alice, 32, Marketing Manager, Sydney]"),
        ("voice", '"Meet Alice. Sunday night. No plan, no groceries, takeaway again. This week is different."'),
    ]),
    ("ACT 1 — Alice Signs Up (0:50 – 1:05)", [
        ("screen", "[MealBuddy login page → Sign up]"),
        ("voice", '"Alice signs up in seconds."'),
        ("action", "Profile modal → fills: Nut-free, Weight loss goal, Shellfish allergy → Save"),
        ("voice", '"MealBuddy now knows exactly who she is — her goals, her allergies, her restrictions. Every answer from here is personalised to her."'),
        ("screen", "[🔥 Day 1 streak appears]"),
        ("voice", '"Day one. The streak has started."'),
    ]),
    ("ACT 2 — Her Documents, Her Context (1:05 – 1:25)", [
        ("screen", "[Files tab]"),
        ("voice", '"Alice has her own personal diet notes — a PDF she never acted on. She uploads it."'),
        ("action", "Upload nutritionist_notes.pdf"),
        ("screen", "[Chat tab]"),
        ("type", '"What are my dietary restrictions based on my uploaded documents?"'),
        ("screen", "[Bot answers accurately from PDF content]"),
        ("voice", '"Not generic advice from the internet. Her own information, read and understood."'),
        ("voice", '"This is RAG — Retrieval-Augmented Generation — powered by Amazon Titan Embeddings V2."'),
    ]),
    ("ACT 3 — Browse & Save Recipes (1:25 – 1:40)", [
        ("screen", "[Chat tab]"),
        ("type", '"Show me high-protein low-carb dinner ideas"'),
        ("screen", "[Bot lists 5 recipes with descriptions and calories]"),
        ("voice", '"MealBuddy searches semantically — not just keywords. It understands what she\'s looking for."'),
        ("type", '"Save the Grilled Chicken Salad to my favourites"'),
        ("screen", "[⭐ Agent action fires — recipe saved]"),
        ("screen", "[Favourites tab — recipe appears]"),
        ("voice", '"One sentence. Saved."'),
    ]),
    ("ACT 4 — One Sentence. Seven Days. (1:40 – 2:05)", [
        ("screen", "[Chat tab]"),
        ("voice", '"Now the moment that used to take Alice an hour every Sunday."'),
        ("type", '"Plan my week with high-protein meals and add to my meal planner"'),
        ("screen", "[Agent Actions panel lights up — 7 × 📅 Add to Meal Plan fires in sequence]"),
        ("voice", '"MealBuddy doesn\'t just answer — it acts. Seven meals planned autonomously using Strands multi-agent orchestration on Amazon Bedrock."'),
        ("screen", "[Switch to Meal Planner tab — fully populated Monday–Sunday]"),
        ("voice", '"Monday through Sunday. Done."'),
        ("type", '"Generate my shopping list for this week"'),
        ("screen", "[🛒 Shopping list populates]"),
        ("screen", "[Switch to Shopping List tab]"),
        ("voice", '"Her grocery run is ready before she\'s stood up from the couch."'),
    ]),
    ("ACT 5 — Tracking That Sticks (2:05 – 2:35)", [
        ("screen", "[Nutrition tab]"),
        ("voice", '"Monday. Alice eats the Grilled Chicken Salad."'),
        ("action", "Log meal → 380 cal → macros fill in → progress bar updates"),
        ("screen", "[Chat tab]"),
        ("type", '"How many calories do I have left today?"'),
        ("screen", "[Bot calls get_nutrition_stats → responds with remaining budget]"),
        ("voice", '"Real-time. Personalised. No manual maths."'),
        ("type", '"Quick high-protein snack under 200 calories?"'),
        ("screen", "[Bot suggests Greek Yogurt Parfait with macros]"),
        ("type", '"How can I adjust the portion sizes to fit my calorie needs?"'),
        ("screen", "[Bot gives personalised portion advice based on her remaining budget and weight loss goal]"),
        ("voice", '"It doesn\'t just suggest food — it coaches her through it. Every answer shaped by her goals, her budget, her day."'),
    ]),
    ("ACT 6 — The Habit Loop (2:35 – 2:45)", [
        ("screen", "[Dashboard — show streak counter and badges panel]"),
        ("voice", '"MealBuddy keeps Alice coming back — not just through results, but through momentum."'),
        ("voice", '"Every day she logs in, her streak grows."'),
        ("voice", '"Hit your first week? Badge unlocked. Log your first meal? Badge unlocked. Each milestone is a small win that builds the habit."'),
        ("voice", '"The app isn\'t just useful. It\'s motivating. Built-in streaks and achievements that make healthy eating stick — not just for a week, but long term."'),
    ]),
    ("CLOSING (2:45 – 3:00)", [
        ("screen", "[Architecture diagram fades in]"),
        ("voice", '"Under the hood: Claude 3 Haiku on Amazon Bedrock for conversation. Titan Embeddings V2 for semantic search. Strands Agents for autonomous multi-step actions. Hosted on Elastic Beanstalk with S3."'),
        ("screen", "[MealBuddy logo + tagline]"),
        ("voice", '"Alice went from no plan to a full week of meals, a shopping list, and real-time calorie tracking — in one conversation with just MealBuddy."'),
        ("voice", '"Your daily buddy for smarter eating — built on AWS."'),
    ]),
]

COLORS = {
    "screen": RGBColor(0x27, 0x6E, 0xB4),   # blue
    "voice":  RGBColor(0x1E, 0x7E, 0x34),   # green
    "action": RGBColor(0xB8, 0x6E, 0x00),   # amber
    "type":   RGBColor(0x6F, 0x42, 0xC1),   # purple
}
LABELS = {"screen": "🎬", "voice": "🎙️", "action": "👆", "type": "⌨️ Type:"}

for heading, items in sections:
    doc.add_heading(heading, level=1)
    for kind, text in items:
        p = doc.add_paragraph()
        label = doc.add_paragraph()
        run = p.add_run(f"{LABELS[kind]}  {text}")
        run.font.color.rgb = COLORS[kind]
        run.font.size = Pt(11)
        if kind == "voice":
            run.font.italic = True
        elif kind == "type":
            run.font.bold = True
    doc.add_paragraph()

# Cheat sheet table
doc.add_heading("📋 Recording Cheat Sheet", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Time'
hdr[1].text = 'Screen'
hdr[2].text = 'Action'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows = [
    ("0:00", "Black screen", "Voiceover only"),
    ("0:15", "3-panel split", "Voiceover only"),
    ("0:50", "Login page", "Sign up, fill profile (Nut-free, Weight loss, Shellfish)"),
    ("1:05", "Files tab", "Upload nutritionist_notes.pdf"),
    ("1:10", "Chat", '"What are my dietary restrictions based on my uploaded documents?"'),
    ("1:25", "Chat", '"Show me high-protein low-carb dinner ideas"'),
    ("1:33", "Chat", '"Save the Grilled Chicken Salad to my favourites"'),
    ("1:36", "Favourites tab", "Show saved recipe"),
    ("1:40", "Chat", '"Plan my week with high-protein meals and add to my meal planner"'),
    ("1:52", "Meal Planner tab", "Show populated week"),
    ("1:56", "Chat", '"Generate my shopping list for this week"'),
    ("2:01", "Shopping List tab", "Show list"),
    ("2:05", "Nutrition tab", "Log Grilled Chicken Salad, 380 cal"),
    ("2:10", "Chat", '"How many calories do I have left today?"'),
    ("2:20", "Chat", '"Quick high-protein snack under 200 calories?"'),
    ("2:25", "Chat", '"How can I adjust the portion sizes to fit my calorie needs?"'),
    ("2:35", "Dashboard", "Show streak counter + badges"),
    ("2:45", "Architecture diagram", "Voiceover only"),
]
for t, s, a in rows:
    row = table.add_row().cells
    row[0].text = t
    row[1].text = s
    row[2].text = a

out = "/mnt/c/Users/Alice.Nguyen/OneDrive - Rio Tinto/Git/awsgenaihackathon/MealBuddy_Demo_Script.docx"
doc.save(out)
print(f"Saved: {out}")
